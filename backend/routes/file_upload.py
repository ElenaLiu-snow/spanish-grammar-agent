from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import os
import tempfile
from pathlib import Path
from typing import Optional
import base64

router = APIRouter()

# 支持的文件类型
SUPPORTED_EXTENSIONS = {
    '.docx': 'word',
    '.md': 'markdown',
    '.txt': 'text',
    '.png': 'image',
    '.jpg': 'image',
    '.jpeg': 'image',
    '.gif': 'image',
    '.webp': 'image'
}


class FileAnalysisRequest(BaseModel):
    """文件分析请求"""
    filename: str
    file_type: str
    content: Optional[str] = None  # 文本内容
    base64_data: Optional[str] = None  # 图片的 base64 数据
    question: Optional[str] = None  # 用户问题（可选）


def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档提取文本和表格，保持原始顺序"""
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph

        doc = Document(file_path)
        result_parts = []

        def iter_block_items(parent):
            """
            生成文档中所有的段落和表格，按照在文档中出现的顺序
            """
            # 确保我们访问正确的 body 元素
            if hasattr(parent, '_element'):
                parent_el = parent._element.body
            else:
                parent_el = parent.element.body

            for child in parent_el.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        # 遍历所有元素
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    result_parts.append(text)

            elif isinstance(block, Table):
                # 处理表格 - 改进的逻辑
                table = block

                # 获取实际的行列数
                rows = table.rows
                if not rows:
                    continue

                # 获取表格的最大列数
                max_cols = max(len(row.cells) for row in rows)

                # 遍历每一行
                for row_idx, row in enumerate(rows):
                    row_data = []

                    # 遍历每个单元格，处理合并单元格
                    for col_idx in range(max_cols):
                        try:
                            # 尝试获取单元格
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                            else:
                                # 处理合并单元格 - 尝试从上一行获取
                                cell = None
                                if row_idx > 0:
                                    prev_row = rows[row_idx - 1]
                                    if col_idx < len(prev_row.cells):
                                        # 检查是否是合并单元格的延伸
                                        cell = prev_row.cells[col_idx]

                            if cell is not None:
                                # 提取单元格中的所有文本
                                cell_text = extract_cell_text(cell)
                                if not cell_text:
                                    cell_text = " "  # 空单元格用空格占位
                                row_data.append(cell_text)
                            else:
                                row_data.append(" ")  # 完全无法获取的单元格

                        except Exception as e:
                            # 单元格处理失败，使用空格占位
                            row_data.append(" ")

                    # 添加这一行到结果
                    markdown_row = '| ' + ' | '.join(row_data) + ' |'
                    result_parts.append(markdown_row)

                    # 在第一行后添加分隔线
                    if row_idx == 0:
                        separator = '|' + '|'.join(['---' for _ in range(max_cols)]) + '|'
                        result_parts.append(separator)

                # 表格结束后添加空行
                result_parts.append('')

        return '\n'.join(result_parts)

    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，请运行: pip install python-docx")
    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=f"读取 Word 文档失败: {str(e)}\n{traceback.format_exc()}")


def extract_cell_text(cell):
    """提取单元格中的所有文本，包括嵌套表格"""
    try:
        # 首先检查单元格中是否有嵌套表格
        if hasattr(cell, '_element'):
            cell_tables = cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl')
            if cell_tables:
                # 有嵌套表格，先提取表格内容
                nested_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        nested_text.append(paragraph.text.strip())
                return ' '.join(nested_text)

        # 正常的单元格文本提取
        cell_paragraphs = [p.text.strip() for p in cell.paragraphs]
        cell_text = ' '.join([p for p in cell_paragraphs if p])
        return cell_text
    except Exception:
        # 如果提取失败，尝试简单方法
        try:
            return cell.text.strip() if hasattr(cell, 'text') else ""
        except:
            return ""


def extract_text_from_markdown(file_path: str) -> str:
    """从 Markdown 文件提取文本"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取 Markdown 文件失败: {str(e)}")


def extract_text_from_text(file_path: str) -> str:
    """从纯文本文件提取内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"读取文本文件失败: {str(e)}")


async def process_file(file: UploadFile) -> dict:
    """处理上传的文件"""
    # 获取文件扩展名
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {file_ext}。支持的类型: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    file_type = SUPPORTED_EXTENSIONS[file_ext]

    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
        content = await file.read()
        tmp_file.write(content)
        tmp_file_path = tmp_file.name

    try:
        result = {
            'filename': file.filename,
            'file_type': file_type,
            'extension': file_ext
        }

        if file_type == 'word':
            result['content'] = extract_text_from_docx(tmp_file_path)

        elif file_type == 'markdown':
            result['content'] = extract_text_from_markdown(tmp_file_path)

        elif file_type == 'text':
            result['content'] = extract_text_from_text(tmp_file_path)

        elif file_type == 'image':
            # 对于图片，返回 base64 编码
            with open(tmp_file_path, 'rb') as img_file:
                image_data = img_file.read()
                result['base64_data'] = base64.b64encode(image_data).decode('utf-8')
                result['size'] = len(image_data)

        return result

    finally:
        # 清理临时文件
        try:
            os.unlink(tmp_file_path)
        except:
            pass


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """上传文件并提取内容"""
    try:
        result = await process_file(file)
        return JSONResponse(content={
            'status': 'success',
            'data': result
        })
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")


@router.post("/analyze")
async def analyze_file(request: FileAnalysisRequest):
    """分析文件内容（调用当前 LLM 模型）"""
    from services.claude import get_claude_service

    try:
        llm_service = get_claude_service()

        # 构建提示词
        if request.file_type == 'image':
            # 图片输入需要模型与服务层支持视觉消息格式
            prompt = request.question or "请描述这张图片中的内容，并识别其中任何与西班牙语语法相关的内容。"

            # TODO: 实现图片分析功能
            return JSONResponse(content={
                'status': 'error',
                'message': '图片分析功能正在开发中，请稍后再试'
            })

        else:
            # 对于文本文档
            content_preview = request.content[:500] if request.content else ""

            if request.question:
                prompt = f"""我有以下文档内容：

```
文件名：{request.filename}
{request.question}

文档内容（节选）：
{content_preview}
```

请根据文档内容回答问题。"""
            else:
                prompt = f"""请分析以下文档内容，并：
1. 总结文档的主要内容
2. 识别其中任何与西班牙语语法相关的内容
3. 如果有语法错误或不规范的表达，请指出并纠正

文件名：{request.filename}

文档内容：
{request.content}"""

            response = await llm_service.chat_complete(prompt)

            return JSONResponse(content={
                'status': 'success',
                'response': response
            })

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"分析失败: {str(e)}")
